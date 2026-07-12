import os
import re
from typing import Optional

import pdfplumber
from docx import Document
from werkzeug.utils import secure_filename

ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}


def validate_resume_upload(file_storage, max_content_length: int | None = None) -> tuple[str, int]:
    """Validate a resume upload and return a safe filename plus file size."""
    if file_storage is None:
        raise ValueError("Please select a file to upload.")

    original_filename = (getattr(file_storage, "filename", "") or "").strip()
    if not original_filename:
        raise ValueError("The selected file is invalid.")

    safe_filename = secure_filename(original_filename)
    if not safe_filename:
        raise ValueError("The selected file is invalid.")

    extension = os.path.splitext(safe_filename)[1].lower()
    if extension not in ALLOWED_RESUME_EXTENSIONS:
        raise ValueError("Only PDF and DOCX files are supported.")

    stream = getattr(file_storage, "stream", None)
    if stream is None:
        raise ValueError("The selected file is invalid.")

    try:
        stream.seek(0, os.SEEK_END)
        file_size = stream.tell()
        stream.seek(0)
    except (AttributeError, OSError) as exc:
        raise ValueError("The selected file is invalid.") from exc

    if file_size <= 0:
        raise ValueError("The selected file is empty.")

    if max_content_length is not None and file_size > max_content_length:
        raise ValueError("File size exceeds the 10 MB limit.")

    return original_filename, file_size


def clean_extracted_text(text: Optional[str]) -> str:
    """Normalize extracted text for storage and display."""
    if not text:
        return ""

    cleaned = re.sub(r"\r\n?", "\n", text)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = "\n".join(line.rstrip() for line in cleaned.splitlines())
    return cleaned.strip()


def extract_resume_text(file_path: str, filename: str) -> str:
    """Extract and clean text from a PDF or DOCX resume file."""
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".pdf":
        return _extract_from_pdf(file_path)
    if extension == ".docx":
        return _extract_from_docx(file_path)

    raise ValueError("Unsupported file format for text extraction.")


def _extract_from_pdf(file_path: str) -> str:
    """Extract all available text from a PDF file."""
    text_parts = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception as exc:  # pragma: no cover - depends on external parser behavior
        raise RuntimeError(f"PDF extraction failed: {exc}") from exc

    if not text_parts:
        raise ValueError("No text could be extracted from the PDF file.")

    return clean_extracted_text("\n\n".join(text_parts))


def _extract_from_docx(file_path: str) -> str:
    """Extract text from paragraphs and tables in a DOCX file."""
    try:
        document = Document(file_path)
    except Exception as exc:  # pragma: no cover - depends on external parser behavior
        raise RuntimeError(f"DOCX extraction failed: {exc}") from exc

    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                text_parts.append(" | ".join(values))

    if not text_parts:
        raise ValueError("No text could be extracted from the DOCX file.")

    return clean_extracted_text("\n".join(text_parts))
